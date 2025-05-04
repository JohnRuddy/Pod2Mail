import asyncio
from flask import Flask, render_template, request, send_file
import requests
import xml.etree.ElementTree as ET
from config import Config
import openai
from docx import Document
import os
import re
import math

app = Flask(__name__)
app.config.from_object(Config)

# Set OpenAI API key
openai.api_key = app.config["OPENAI_API_KEY"]

@app.route("/")
def home():
    return render_template("home.html")

@app.route("/about")
def about():
    return render_template("about.html")

def get_audio_url(item):
    """
    Extract the audio file URL from an RSS feed item.
    Supports MP3, MP4, and WAV formats.
    """
    # Check the <enclosure> tag
    enclosure = item.find("enclosure")
    if enclosure is not None and "url" in enclosure.attrib:
        url = enclosure.get("url")
        if url.endswith((".mp3", ".mp4", ".wav")):
            return url

    # Check other fields (e.g., <link>)
    link = item.find("link")
    if link is not None and link.text.endswith((".mp3", ".mp4", ".wav")):
        return link.text

    # Add more checks if needed for other fields
    return None

@app.route("/load_podcast", methods=["GET", "POST"])
def load_podcast():
    episodes = []
    podcast_xml = None  # Variable to store the XML content for inspection
    if request.method == "POST":
        rss_url = request.form.get("rss_url")
        try:
            response = requests.get(rss_url)
            response.raise_for_status()  # Raise an error for bad responses
            podcast_xml = response.content.decode("utf-8")  # Decode and store the XML
            print(f"Podcast XML:\n{podcast_xml}")  # Debug statement to inspect the XML

            root = ET.fromstring(response.content)

            # Parse the RSS feed
            for item in root.findall(".//item")[:10]:  # Get up to 10 episodes
                title = item.find("title").text if item.find("title") is not None else "No Title"
                link = item.find("link").text if item.find("link") is not None else "No Link"
                pub_date = item.find("pubDate").text if item.find("pubDate") is not None else "Unknown"
                audio_url = get_audio_url(item)  # Use the helper function to find the audio URL
                episodes.append({
                    "id": link,
                    "title": title, 
                    "published_date": pub_date,
                    "audio_url": audio_url,
                })
        except Exception as e:
            print(f"Error fetching or parsing RSS feed: {e}")
    return render_template("load_podcast.html", episodes=episodes)

@app.route("/summarize", methods=["POST"])
async def summarize():
    selected_episodes = request.form.getlist("selected_episodes")
    episodes = request.form.to_dict(flat=False)  # Get all form data as a dictionary
    summaries = []

    async def process_episode(episode_id, audio_url):
        try:
            # Sanitize the episode_id to create a valid filename
            sanitized_episode_id = re.sub(r'[^\w\-]', '_', episode_id)  # Replace invalid characters with underscores
            audio_file_path = f"episode_{sanitized_episode_id}.mp3"

            # Fetch the audio content
            print(f"Fetching audio file for episode {episode_id} from URL: {audio_url}")  # Debug statement
            response = requests.get(audio_url)
            response.raise_for_status()  # Ensure the request was successful
            audio_content = response.content

            # Save the audio file locally
            with open(audio_file_path, "wb") as f:
                f.write(audio_content)
            print(f"Audio file saved to {audio_file_path}")  # Debug statement

            # Check file size and split if necessary
            max_size = 25 * 1024 * 1024  # 25 MB limit
            file_size = os.path.getsize(audio_file_path)
            if file_size > max_size:
                print(f"File size ({file_size} bytes) exceeds limit. Splitting file...")
                audio_chunks = split_file(audio_file_path, max_size)
            else:
                audio_chunks = [audio_file_path]

            # Process each chunk
            transcription_text = ""
            for chunk_path in audio_chunks:
                chunk_size = os.path.getsize(chunk_path)
                if chunk_size > max_size:
                    print(f"Chunk size ({chunk_size} bytes) still exceeds limit. Skipping chunk: {chunk_path}")
                    continue  # Skip chunks that still exceed the limit

                print(f"Processing chunk: {chunk_path}")
                with open(chunk_path, "rb") as chunk_file:
                    transcription_response = await openai.Audio.atranscribe(
                        model="whisper-1",
                        file=chunk_file
                    )
                    transcription_text += transcription_response["text"] + "\n"

            # Split transcription text into smaller chunks for summarization
            max_tokens = 8000  # Leave some buffer for the prompt
            transcription_chunks = split_text_into_chunks(transcription_text, max_tokens)

            # Summarize each chunk
            summary_text = ""
            for i, chunk in enumerate(transcription_chunks):
                print(f"Summarizing chunk {i + 1}/{len(transcription_chunks)}")
                summary_response = await openai.ChatCompletion.acreate(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are a summarization assistant."},
                        {"role": "user", "content": f"Summarize the following text into bullet points:\n\n{chunk}"},
                    ]
                )
                summary_text += summary_response["choices"][0]["message"]["content"] + "\n"

            # Save the summary to a Word document
            summary_file_path = f"summary_{sanitized_episode_id}.docx"
            doc = Document()
            doc.add_heading(f"Summary for Episode {episode_id}", level=1)
            for bullet in summary_text.split("\n"):
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style="List Bullet")
            doc.save(summary_file_path)

            # Clean up temporary chunk files
            for chunk_path in audio_chunks:
                os.remove(chunk_path)

            return {"episode_id": episode_id, "file_path": summary_file_path, "summary": summary_text.strip()}
        except Exception as e:
            print(f"Error processing episode {episode_id}: {e}")
            return None

    def split_file(file_path, max_size):
        """
        Splits a file into smaller chunks if it exceeds the max_size.
        Returns a list of chunk file paths.
        """
        chunk_paths = []
        base_name, ext = os.path.splitext(file_path)  # Separate the base name and extension
        with open(file_path, "rb") as f:
            chunk_index = 0
            while True:
                # Read slightly less than max_size to account for overhead
                chunk_data = f.read(max_size - 1024)
                if not chunk_data:
                    break
                # Create a chunk file with the same extension
                chunk_path = f"{base_name}_chunk{chunk_index}{ext}"
                with open(chunk_path, "wb") as chunk_file:
                    chunk_file.write(chunk_data)
                chunk_paths.append(chunk_path)
                chunk_index += 1
        return chunk_paths

    def split_text_into_chunks(text, max_tokens):
        """
        Splits a large text into smaller chunks based on the maximum token limit.
        """
        words = text.split()
        chunks = []
        current_chunk = []

        for word in words:
            current_chunk.append(word)
            # Estimate token count (1 word ≈ 1 token for simplicity)
            if len(current_chunk) >= max_tokens:
                chunks.append(" ".join(current_chunk))
                current_chunk = []

        # Add the last chunk if it exists
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    # Process all selected episodes concurrently
    tasks = []
    for episode_id in selected_episodes:
        # Find the episode's audio URL
        audio_url = None
        for i in range(len(episodes.get("episodes[0][id]", []))):
            if episodes[f"episodes[{i}][id]"][0] == episode_id:
                audio_url = episodes[f"episodes[{i}][audio_url]"][0]
                break

        if audio_url:
            tasks.append(process_episode(episode_id, audio_url))

    # Wait for all tasks to complete
    results = await asyncio.gather(*tasks)

    # Filter out any failed tasks
    summaries = [result for result in results if result is not None]

    return render_template("summarize.html", summaries=summaries)

@app.route("/download/<path:filename>")
def download_file(filename):
    return send_file(filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
